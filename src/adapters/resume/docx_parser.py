"""DocxResumeParser — extract resume text from .docx bytes (ResumeParserPort).

The sibling of ``PyPDF2ResumeParser`` for Word documents: the browser upload (W5)
accepts ``.docx`` as well as ``.pdf``, and both formats reach ``ResumeService``
through the same ``ResumeParserPort``. python-docx is already a dependency (F uses
it on the write side); keeping it behind the port keeps the library out of the core.
"""

import io
from zipfile import BadZipFile

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from src.core.ports.resume_parser_port import ResumeParserPort


class DocxResumeParser(ResumeParserPort):
    """Extract plain text from a ``.docx`` resume using python-docx."""

    def extract_text(self, data: bytes) -> str:
        """Extract whitespace-stripped text from the given ``.docx`` bytes.

        Pulls text from body paragraphs and table cells (skills are often laid out
        in tables), joins them with newlines, and strips.

        Args:
            data: The raw ``.docx`` file contents.

        Returns:
            The concatenated, stripped document text.

        Raises:
            ValueError: When the bytes are not a valid ``.docx`` package or yield
                no extractable text.
        """
        try:
            document = Document(io.BytesIO(data))
        except (PackageNotFoundError, BadZipFile, KeyError) as exc:
            raise ValueError("The uploaded file is not a valid .docx document") from exc

        lines = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    lines.append(cell.text)

        raw_text = "\n".join(line for line in lines if line).strip()

        if not raw_text:
            raise ValueError("No text could be extracted from the resume .docx")

        return raw_text
