"""DocxWriter — renders generated documents to ``.docx`` via python-docx (F).

The one adapter implementing ``DocxWriterPort`` for both artifact types (the F1+F2
merge of ADR-029's writer). It renders the **already-formatted** structured document
— the deterministic formatter has run first — so its only jobs are layout and
guaranteeing the ``•`` bullet glyph (CLAUDE.md #6). The file is written for every
outcome, including ``needs_review`` (with the formatter's review markers intact), so
a generation is never lost to a formatting nit (ADR-029). Keeping python-docx behind
the port keeps the rendering library out of the core.
"""

import os
import re

from docx import Document
from docx.document import Document as DocxDocument

from src.core.domain.cover_letter import CoverLetter
from src.core.domain.tailored_resume import TailoredResume
from src.core.ports.docx_writer_port import DocxWriterPort

_LEADING_BULLET = re.compile(r"^[•\-*]\s*")


class DocxWriter(DocxWriterPort):
    """Render tailored resumes and cover letters to ``.docx`` files."""

    def write_resume(self, doc: TailoredResume, path: str) -> None:
        """Render a tailored resume (summary, skills, sections) to ``path``."""
        document = Document()
        document.add_heading("Professional Summary", level=1)
        document.add_paragraph(doc.summary)

        if doc.skills:
            document.add_heading("Skills", level=1)
            document.add_paragraph(", ".join(doc.skills))

        for section in doc.sections:
            document.add_heading(section.heading, level=1)
            for bullet in section.bullets:
                document.add_paragraph(self._bullet(bullet))

        self._save(document, path)

    def write_cover_letter(self, doc: CoverLetter, path: str) -> None:
        """Render a cover letter (salutation, paragraphs, closing) to ``path``."""
        document = Document()
        document.add_paragraph(doc.salutation)
        for paragraph in doc.paragraphs:
            document.add_paragraph(paragraph)
        document.add_paragraph(doc.closing)
        self._save(document, path)

    @staticmethod
    def _bullet(text: str) -> str:
        """Return ``text`` as a single ``•``-prefixed bullet line (no double marker).

        Args:
            text: The bullet's text (which may already carry a leading marker).

        Returns:
            The text prefixed with exactly one ``•`` bullet glyph.
        """
        return "• " + _LEADING_BULLET.sub("", text.strip())

    @staticmethod
    def _save(document: DocxDocument, path: str) -> None:
        """Create the parent directory if needed and save the document to ``path``.

        Args:
            document: The python-docx document to persist.
            path: Destination ``.docx`` path.
        """
        parent = os.path.dirname(path)
        if parent:
            os.makedirs(parent, exist_ok=True)
        document.save(path)
