"""Unit tests for DocxResumeParser.

Uses python-docx (already a dependency) to build real ``.docx`` bytes in memory —
a round-trip through the library, no real file I/O and no network.
"""

import io

import pytest
from docx import Document

from src.adapters.resume.docx_parser import DocxResumeParser


def _docx_bytes(paragraphs: list[str] = (), table_rows: list[list[str]] = ()) -> bytes:
    """Return the bytes of a ``.docx`` with the given paragraphs and table rows."""
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for r, row in enumerate(table_rows):
            for c, cell_text in enumerate(row):
                table.cell(r, c).text = cell_text
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_extract_text_reads_paragraphs():
    """Body paragraphs are extracted and joined."""
    data = _docx_bytes(paragraphs=["Senior Engineer", "8 years experience"])
    text = DocxResumeParser().extract_text(data)
    assert "Senior Engineer" in text
    assert "8 years experience" in text


def test_extract_text_reads_table_cells():
    """Skills laid out in a table are extracted from the cells."""
    data = _docx_bytes(paragraphs=["Skills"], table_rows=[["Python", "FastAPI"]])
    text = DocxResumeParser().extract_text(data)
    assert "Python" in text
    assert "FastAPI" in text


def test_extract_text_empty_docx_raises_value_error():
    """A document with no text yields a ValueError."""
    data = _docx_bytes(paragraphs=["", "   "])
    with pytest.raises(ValueError, match="No text could be extracted"):
        DocxResumeParser().extract_text(data)


def test_extract_text_non_docx_bytes_raises_value_error():
    """Bytes that are not a valid .docx package raise a clear ValueError."""
    with pytest.raises(ValueError, match="not a valid .docx"):
        DocxResumeParser().extract_text(b"not a real docx file")
